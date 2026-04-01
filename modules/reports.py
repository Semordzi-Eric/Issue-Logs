import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib import colors
import io
import datetime
from database.models import get_session, Issue, EmailLog, EmailResponse


# ──────────────────────────────────────────────
# DATA HELPERS
# ──────────────────────────────────────────────
def _load_full_data():
    session = get_session()
    issues  = session.query(Issue).all()
    emails  = session.query(EmailLog).all()

    issue_data = []
    for i in issues:
        i_emails = [e for e in emails if e.issue_id == i.id]
        responded = sum(1 for e in i_emails if e.response_status == "Responded")
        pending   = sum(1 for e in i_emails if e.response_status != "Responded")
        issue_data.append({
            'Issue ID':       i.issue_id,
            'Date':           pd.to_datetime(i.date),
            'Title':          i.title,
            'Description':    i.description or "",
            'Category':       i.category,
            'Priority':       i.priority,
            'Affected System':i.affected_system,
            'Transaction ID': i.transaction_id or "-",
            'Amount':         i.amount or 0.0,
            'Root Cause':     i.root_cause or "Not yet determined",
            'Status':         i.status,
            'Resolution':     i.resolution_notes or "Pending",
            'Emails Sent':    len(i_emails),
            'Responded':      responded,
            'Pending Emails': pending,
        })

    session.close()
    return pd.DataFrame(issue_data)


# ──────────────────────────────────────────────
# WORD REPORT
# ──────────────────────────────────────────────
def generate_word_report(df: pd.DataFrame, start_date: datetime.date, end_date: datetime.date) -> bytes:
    doc = Document()
    period_str = f"{start_date} to {end_date}"

    # Title
    title = doc.add_heading(f'Audit & Revenue Assurance Report', 0)
    sub   = doc.add_paragraph(f'Period: {period_str}    |    Generated: {datetime.date.today()}')
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x60, 0x78)
    doc.add_paragraph()

    total_issues   = len(df)
    resolved       = len(df[df['Status'] == 'Resolved'])
    open_count     = len(df[df['Status'] == 'Open'])
    investigating  = len(df[df['Status'] == 'Investigating'])
    escalated      = len(df[df['Status'] == 'Escalated'])
    total_amount   = df['Amount'].sum()
    high_count     = len(df[df['Priority'].isin(['High', 'Critical'])])
    total_emails   = df['Emails Sent'].sum()
    pending_emails = df['Pending Emails'].sum()

    # ── 1. Executive Summary ──────────────────────
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        f"During the period of {period_str}, the Audit & Revenue Assurance team logged a total of "
        f"{total_issues} issue(s). Of these, {resolved} have been fully resolved, {investigating} are "
        f"under active investigation, {escalated} have been escalated, and {open_count} remain open. "
        f"The total estimated revenue impact for this period stands at GHS ₵{total_amount:,.2f}. "
        f"A total of {high_count} high or critical priority incidents were recorded, requiring immediate attention. "
        f"{total_emails} email communication(s) were dispatched during this period, of which "
        f"{pending_emails} are still awaiting a response."
    )

    # ── 2. Issue Summary Table ────────────────────
    doc.add_heading('2. Issue Summary', level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Issue ID', 'Title', 'Category', 'Priority', 'Status', 'Amount (GHS ₵)']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    for _, row in df.iterrows():
        r = table.add_row().cells
        r[0].text = row['Issue ID']
        r[1].text = row['Title']
        r[2].text = row['Category']
        r[3].text = row['Priority']
        r[4].text = row['Status']
        r[5].text = f"₵{row['Amount']:,.2f}"

    doc.add_paragraph()

    # ── 3. Breakdown by Category ──────────────────
    doc.add_heading('3. Breakdown by Category', level=1)
    cat_counts = df['Category'].value_counts()
    for cat, count in cat_counts.items():
        pct = (count / total_issues) * 100
        doc.add_paragraph(f"{cat}: {count} issue(s)  ({pct:.1f}%)", style='List Bullet')

    # ── 4. Key Incidents ──────────────────────────
    doc.add_heading('4. Key Incidents (High & Critical Priority)', level=1)
    high_df = df[df['Priority'].isin(['High', 'Critical'])]
    if high_df.empty:
        doc.add_paragraph("No high or critical priority incidents were reported this month.")
    else:
        for _, row in high_df.iterrows():
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(f"[{row['Issue ID']}] {row['Title']} ({row['Priority']})")
            r.bold = True
            doc.add_paragraph(
                f"   System: {row['Affected System']}  |  Category: {row['Category']}  |  "
                f"Status: {row['Status']}  |  Amount: GHS ₵{row['Amount']:,.2f}"
            )
            doc.add_paragraph(f"   Description: {row['Description']}")
            doc.add_paragraph(f"   Root Cause: {row['Root Cause']}")
            doc.add_paragraph(f"   Resolution: {row['Resolution']}")
            doc.add_paragraph()

    # ── 5. Revenue Impact Analysis ────────────────
    doc.add_heading('5. Revenue Impact Analysis', level=1)
    doc.add_paragraph(f"Total Estimated Impact: GHS ₵{total_amount:,.2f}")
    by_category = df.groupby('Category')['Amount'].sum().sort_values(ascending=False)
    for cat, amt in by_category.items():
        if amt > 0:
            doc.add_paragraph(f"{cat}: GHS ₵{amt:,.2f}", style='List Bullet')

    # ── 6. Communications Summary ─────────────────
    doc.add_heading('6. Email Communications Summary', level=1)
    doc.add_paragraph(
        f"Total emails sent: {int(total_emails)}.  "
        f"Awaiting response: {int(pending_emails)}.  "
        f"Response rate: {((total_emails - pending_emails) / total_emails * 100):.1f}% "
        if total_emails > 0 else "No emails were logged this period."
    )

    # ── 7. Outstanding Issues ─────────────────────
    doc.add_heading('7. Outstanding Issues (Unresolved)', level=1)
    open_df = df[df['Status'] != 'Resolved']
    if open_df.empty:
        doc.add_paragraph("All issues for this period have been resolved. Excellent work.")
    else:
        for _, row in open_df.iterrows():
            doc.add_paragraph(
                f"[{row['Issue ID']}] {row['Title']} — Status: {row['Status']}  |  "
                f"Priority: {row['Priority']}",
                style='List Number'
            )

    # ── 8. Recommendations ───────────────────────
    doc.add_heading('8. Recommendations', level=1)
    recs = [
        "Ensure all 'Open' issues are assigned to a named owner with a target resolution date.",
        "Escalate critical and high-priority issues that have been unresolved for more than 5 business days.",
        "Follow up on all emails with no response within 48 hours of the original send date.",
        "Conduct a root cause analysis for repeated categories (e.g. reversals, failed transactions).",
        "Review affected systems with the highest incident counts for systemic fixes.",
    ]
    for rec in recs:
        doc.add_paragraph(rec, style='List Bullet')

    doc.add_heading('9. Disclaimer', level=1)
    doc.add_paragraph(
        "This report is generated from internal audit logs and is intended for authorised personnel only. "
        f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} by the "
        "Audit & Revenue Assurance Tracker system."
    )

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ──────────────────────────────────────────────
# PDF REPORT
# ──────────────────────────────────────────────
def generate_pdf_report(df: pd.DataFrame, start_date: datetime.date, end_date: datetime.date) -> bytes:
    bio    = io.BytesIO()
    period_str = f"{start_date} to {end_date}"
    doc    = SimpleDocTemplate(bio, pagesize=A4,
                               rightMargin=2*cm, leftMargin=2*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()

    TITLE_STYLE  = ParagraphStyle('title',  fontSize=18, fontName='Helvetica-Bold',
                                  spaceAfter=4, textColor=colors.HexColor('#1A1D2E'))
    HEAD1_STYLE  = ParagraphStyle('h1',     fontSize=13, fontName='Helvetica-Bold',
                                  spaceBefore=12, spaceAfter=4,
                                  textColor=colors.HexColor('#2E7D32'))
    BODY_STYLE   = ParagraphStyle('body',   fontSize=9,  fontName='Helvetica',
                                  leading=14, spaceAfter=4)
    BULLET_STYLE = ParagraphStyle('bullet', fontSize=9,  fontName='Helvetica',
                                  leading=14, leftIndent=12, spaceAfter=3,
                                  bulletIndent=0, bulletText='•')
    META_STYLE   = ParagraphStyle('meta',   fontSize=8,  fontName='Helvetica',
                                  textColor=colors.grey)

    total_issues  = len(df)
    resolved      = len(df[df['Status'] == 'Resolved'])
    open_count    = len(df[df['Status'] == 'Open'])
    investigating = len(df[df['Status'] == 'Investigating'])
    escalated     = len(df[df['Status'] == 'Escalated'])
    total_amount  = df['Amount'].sum()
    high_count    = len(df[df['Priority'].isin(['High', 'Critical'])])
    total_emails  = int(df['Emails Sent'].sum())
    pending_emails = int(df['Pending Emails'].sum())

    story = []

    # ── Cover ─────────────────────────────────────
    story.append(Paragraph("Audit & Revenue Assurance", TITLE_STYLE))
    story.append(Paragraph(f"Period: {period_str}", HEAD1_STYLE))
    story.append(Paragraph(f"Generated: {datetime.datetime.now().strftime('%d %B %Y %H:%M')}",
                            META_STYLE))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#2E7D32')))
    story.append(Spacer(1, 0.3*cm))

    # ── 1. Executive Summary ──────────────────────
    story.append(Paragraph("1. Executive Summary", HEAD1_STYLE))
    story.append(Paragraph(
        f"During the period {period_str}, the team logged <b>{total_issues}</b> issue(s). "
        f"<b>{resolved}</b> resolved, <b>{investigating}</b> under investigation, "
        f"<b>{escalated}</b> escalated, <b>{open_count}</b> open. "
        f"Total estimated revenue impact: <b>GHS ₵{total_amount:,.2f}</b>. "
        f"<b>{high_count}</b> high/critical incidents recorded. "
        f"<b>{total_emails}</b> emails sent; <b>{pending_emails}</b> awaiting response.",
        BODY_STYLE
    ))
    story.append(Spacer(1, 0.3*cm))

    # ── 2. Issue Summary Table ────────────────────
    story.append(Paragraph("2. Issue Summary", HEAD1_STYLE))
    tbl_data = [['Issue ID', 'Title', 'Category', 'Priority', 'Status', 'Amount (₵)']]
    for _, row in df.iterrows():
        tbl_data.append([
            row['Issue ID'],
            (row['Title'][:30] + '...') if len(row['Title']) > 30 else row['Title'],
            row['Category'],
            row['Priority'],
            row['Status'],
            f"₵{row['Amount']:,.0f}"
        ])
    tbl = Table(tbl_data, repeatRows=1,
                colWidths=[2.8*cm, 5*cm, 3.5*cm, 2*cm, 2.5*cm, 2.5*cm])
    tbl.setStyle(TableStyle([
        ('BACKGROUND',   (0, 0), (-1, 0),  colors.HexColor('#2E7D32')),
        ('TEXTCOLOR',    (0, 0), (-1, 0),  colors.white),
        ('FONTNAME',     (0, 0), (-1, 0),  'Helvetica-Bold'),
        ('FONTSIZE',     (0, 0), (-1, -1), 8),
        ('ROWBACKGROUNDS',(0, 1), (-1, -1), [colors.HexColor('#F5F7FA'), colors.white]),
        ('GRID',         (0, 0), (-1, -1), 0.4, colors.HexColor('#D0D6EE')),
        ('VALIGN',       (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING',  (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING',   (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING',(0, 0), (-1, -1), 3),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.4*cm))

    # ── 3. Category Breakdown ─────────────────────
    story.append(Paragraph("3. Breakdown by Category", HEAD1_STYLE))
    for cat, count in df['Category'].value_counts().items():
        pct = (count / total_issues) * 100
        story.append(Paragraph(f"{cat}: {count} issue(s)  ({pct:.1f}%)", BULLET_STYLE))
    story.append(Spacer(1, 0.3*cm))

    # ── 4. Key Incidents ──────────────────────────
    story.append(Paragraph("4. Key Incidents (High & Critical)", HEAD1_STYLE))
    high_df = df[df['Priority'].isin(['High', 'Critical'])]
    if high_df.empty:
        story.append(Paragraph("No high or critical incidents this month.", BODY_STYLE))
    else:
        for _, row in high_df.iterrows():
            story.append(Paragraph(
                f"<b>[{row['Issue ID']}] {row['Title']}</b> — Priority: {row['Priority']} | "
                f"Status: {row['Status']} | System: {row['Affected System']} | "
                f"Amount: GHS ₵{row['Amount']:,.2f}",
                BULLET_STYLE
            ))
            story.append(Paragraph(
                f"&nbsp;&nbsp;Root Cause: {row['Root Cause']}", BODY_STYLE))
            story.append(Paragraph(
                f"&nbsp;&nbsp;Resolution: {row['Resolution']}", BODY_STYLE))
    story.append(Spacer(1, 0.3*cm))

    # ── 5. Revenue Impact ─────────────────────────
    story.append(Paragraph("5. Revenue Impact Analysis", HEAD1_STYLE))
    story.append(Paragraph(f"<b>Total:</b> GHS ₵{total_amount:,.2f}", BODY_STYLE))
    by_cat = df.groupby('Category')['Amount'].sum().sort_values(ascending=False)
    for cat, amt in by_cat.items():
        if amt > 0:
            story.append(Paragraph(f"{cat}: GHS ₵{amt:,.2f}", BULLET_STYLE))
    story.append(Spacer(1, 0.3*cm))

    # ── 6. Outstanding Issues ─────────────────────
    story.append(Paragraph("6. Outstanding Issues", HEAD1_STYLE))
    open_df = df[df['Status'] != 'Resolved']
    if open_df.empty:
        story.append(Paragraph("All issues resolved for this period.", BODY_STYLE))
    else:
        for _, row in open_df.iterrows():
            story.append(Paragraph(
                f"[{row['Issue ID']}] {row['Title']} — {row['Status']} | {row['Priority']}",
                BULLET_STYLE
            ))
    story.append(Spacer(1, 0.3*cm))

    # ── 7. Recommendations ───────────────────────
    story.append(Paragraph("7. Recommendations", HEAD1_STYLE))
    for rec in [
        "Assign named owners to all open issues with target resolution dates.",
        "Escalate unresolved critical issues older than 5 business days.",
        "Follow up on emails with no response within 48 hours.",
        "Review systems with highest incident rates for systemic fixes.",
        "Conduct monthly root cause analysis for repeat categories.",
    ]:
        story.append(Paragraph(rec, BULLET_STYLE))
    story.append(Spacer(1, 0.3*cm))

    # ── Footer ────────────────────────────────────
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    story.append(Paragraph(
        f"Confidential — Audit & Revenue Assurance Tracker | {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}",
        META_STYLE
    ))

    doc.build(story)
    return bio.getvalue()


# ──────────────────────────────────────────────
# RENDER
# ──────────────────────────────────────────────
def render_reports():
    st.title("📄 Monthly Audit Report Generator")
    st.markdown("Generate detailed, export-ready audit reports for any logged month.")
    st.markdown("---")

    df = _load_full_data()

    if df.empty:
        st.info("No audit data available. Log issues first.")
        return

    st.subheader("📅 Select Reporting Period")
    c1, c2 = st.columns([2, 3])
    with c1:
        # Default to last 30 days
        today = datetime.date.today()
        month_ago = today - datetime.timedelta(days=30)
        
        date_range = st.date_input(
            "Select Range (Start to End)",
            value=(month_ago, today),
            help="Click once to select Start date, then once for End date."
        )

    # Only filter if a full range is selected
    if isinstance(date_range, (list, tuple)) and len(date_range) == 2:
        start_date, end_date = date_range
        mask = (df['Date'].dt.date >= start_date) & (df['Date'].dt.date <= end_date)
        range_df = df[mask]
        selected_period_label = f"{start_date} to {end_date}"
    else:
        st.warning("Please select both a Start and End date in the calendar.")
        return

    # ── Preview KPIs ─────────────────────────────
    st.markdown(f"#### 📊 Preview — {selected_period_label}")
    p1, p2, p3, p4, p5 = st.columns(5)
    p1.metric("Total Issues",     len(range_df))
    p2.metric("Resolved",         len(range_df[range_df['Status'] == 'Resolved']))
    p3.metric("High / Critical",  len(range_df[range_df['Priority'].isin(['High', 'Critical'])]))
    p4.metric("Revenue Impact",   f"₵{range_df['Amount'].sum():,.2f}")
    p5.metric("Emails Sent",      int(range_df['Emails Sent'].sum()))

    if range_df.empty:
        st.info(f"No audit logs found for the selected range: {selected_period_label}")
        return

    # ── Issue detail table ───────────────────────
    with st.expander("📋 View Issue Data for this Period", expanded=False):
        st.dataframe(
            range_df[['Issue ID','Title','Category','Priority','Status','Amount','Emails Sent','Resolution']],
            hide_index=True,
            use_container_width=True
        )

    st.markdown("---")
    st.markdown("### 📥 Export Report")
    st.markdown("Full report sections including Executive Summary, Key Incidents, Revenue Analysis, and Recommendations.")

    with st.spinner("Preparing reports..."):
        docx_data = generate_word_report(range_df, start_date, end_date)
        pdf_data  = generate_pdf_report(range_df, start_date, end_date)
        csv_data  = range_df.to_csv(index=False).encode('utf-8')

    exp_col1, exp_col2, exp_col3 = st.columns(3)

    with exp_col1:
        st.download_button(
            label="📄 Word (DOCX)",
            data=docx_data,
            file_name=f"Audit_Report_{start_date}_to_{end_date}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )

    with exp_col2:
        st.download_button(
            label="📕 PDF Report",
            data=pdf_data,
            file_name=f"Audit_Report_{start_date}_to_{end_date}.pdf",
            mime="application/pdf"
        )

    with exp_col3:
        st.download_button(
            label="📊 Raw Data (CSV)",
            data=csv_data,
            file_name=f"Audit_Data_{start_date}_to_{end_date}.csv",
            mime="text/csv"
        )
